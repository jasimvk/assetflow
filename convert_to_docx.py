#!/usr/bin/env python3
"""
Convert CLIENT_README.md to DOCX format
Requires: pip install python-docx markdown2
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    import sys
except ImportError:
    print("Installing required packages...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to formatted DOCX"""
    
    # Create document
    doc = Document()
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split into lines
    lines = content.split('\n')
    
    in_code_block = False
    code_lines = []
    
    for line in lines:
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block - add collected lines
                if code_lines:
                    p = doc.add_paragraph('\n'.join(code_lines))
                    p.style = 'Intense Quote'
                    code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Skip empty lines
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # Handle horizontal rules
        elif line.startswith('---'):
            p = doc.add_paragraph('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Handle bullet lists
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('+ '):
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line[2:])  # Remove bold markers
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Remove italic markers
            text = re.sub(r'`(.*?)`', r'\1', text)  # Remove code markers
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Remove links
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
        
        # Handle checkboxes
        elif '- [ ]' in line or '- [x]' in line or '- [X]' in line:
            checked = '[x]' in line.lower()
            text = re.sub(r'- \[[ xX]\]\s*', '', line)
            text = ('☑ ' if checked else '☐ ') + text
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Handle tables
        elif '|' in line and line.strip().startswith('|'):
            # Simple table handling - could be enhanced
            continue
        
        # Regular paragraphs
        else:
            text = line
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'"\1"', text)  # Code
            text = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1 (\2)', text)  # Links
            
            if text.strip():
                p = doc.add_paragraph(text)
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Successfully converted {md_file} to {docx_file}")
    print(f"📄 Open the file to view the formatted document")

if __name__ == '__main__':
    input_file = 'CLIENT_README.md'
    output_file = 'AssetFlow_User_Guide.docx'
    
    try:
        parse_markdown_to_docx(input_file, output_file)
    except FileNotFoundError:
        print(f"❌ Error: {input_file} not found!")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Error: {e}")
        sys.exit(1)
